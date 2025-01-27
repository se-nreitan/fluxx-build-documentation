# HOW TO USE:
# 1. Install dependencies: pip install -r requirements.txt
# 2. Make sure you have Google Chrome installed
# 3. Run this script: python "Fluxx Build Documentation Data Scraper.py"

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import WebDriverException, TimeoutException
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from getpass import getpass
from urllib.parse import urlparse
import sys
import platform
import os
import re
import subprocess
import json
import zipfile
import requests
import winreg
import shutil
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import threading

print("Script starting...")
print(f"Python version: {sys.version}")
print(f"Current working directory: {os.getcwd()}")

def print_logo():
    """Print the Social Edge logo and contact info"""
    logo = """
                    ;X.
                  ;XX+
                .XX+  ..:.
              .XX+  xXXxxXXX
            .XXx  xXX.    .XX
            XX  ;XX. .XX+  XX
            Xx  .  .XXX  +XX.
            ;XXx.;XXx  ;XX;
              .+++:  :XX;
                   :XX+
                  +X+

       Social Edge Consulting - 2025
   nick.reitan@socialedgeconsulting.com


  Fluxx Build Documentation Automated Tool

==================================================
"""
    print(logo)

def get_chrome_path():
    """Get installed Chrome path from registry"""
    try:
        # Check common Chrome locations
        possible_paths = [
            r"C:\Program Files\Google\Chrome\Application\chrome.exe",
            r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        ]
        
        # Try registry first
        try:
            key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\chrome.exe")
            chrome_path = winreg.QueryValue(key, None)
            if os.path.exists(chrome_path):
                return chrome_path
        except:
            pass
        
        # Try common paths
        for path in possible_paths:
            if os.path.exists(path):
                return path
                
        return None
    except:
        return None

def get_resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def get_fluxx_url():
    """Get and validate Fluxx URL from user input"""
    while True:
        url = input("\nEnter Fluxx Instance Name or URL (e.g., 'example' or example.fluxx.io): ").strip()
        if not url:
            print("URL cannot be empty. Please try again.")
            continue
                
        url = validate_fluxx_url(url)
        print(f"\nUsing URL: {url}")
        verify = input("Is this correct? (y/n): ").strip().lower()
        if verify == 'y':
            return url
        print()  # Add blank line before retry

def get_chrome_version():
    """Get Chrome version from registry."""
    try:
        import winreg
        key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r'Software\Google\Chrome\BLBeacon')
        version = winreg.QueryValueEx(key, 'version')[0]
        return version
    except:
        return "unknown"

def get_driver_version(driver_path):
    """Get ChromeDriver version by running the executable with --version flag."""
    try:
        result = subprocess.run([driver_path, '--version'], 
                              capture_output=True, 
                              text=True)
        version = result.stdout.split()[1]  # Format: "ChromeDriver XX.X.XXXX.XX"
        return version.split('.')[0]  # Return major version
    except:
        return None

def setup_webdriver():
    try:
        print("\nSetting up Chrome WebDriver...")
        
        # Use bundled ChromeDriver
        driver_path = get_resource_path("chromedriver.exe")
        chrome_binary_path = get_resource_path(os.path.join("chrome-portable", "chrome.exe"))
        
        print(f"Looking for Chrome at: {chrome_binary_path}")
        
        # Basic Chrome setup
        options = webdriver.ChromeOptions()
        options.binary_location = chrome_binary_path
        options.add_argument('--no-sandbox')  # Required for navigation
        options.add_argument('--disable-dev-shm-usage')  # Required for navigation
        
        print("Initializing Chrome...")
        service = Service(driver_path)
        driver = webdriver.Chrome(service=service, options=options)
        
        # Test navigation to make sure it works
        print("Testing navigation...")
        driver.get("https://www.google.com")
        if driver.current_url == "data:,":
            raise Exception("Chrome navigation not working properly")
            
        print("Chrome initialized successfully!")
        return driver
        
    except Exception as e:
        print("\nError setting up Chrome WebDriver:")
        print(str(e))
        print("\nChrome path:", chrome_binary_path)
        print("Driver path:", driver_path)
        input("\nPress Enter to exit...")
        raise

def get_credentials():
    """Get username and password from user securely."""
    print("\nPlease enter your Fluxx credentials:")
    username = input("Username: ").strip()
    password = getpass("Password: ")  # getpass hides the password while typing
    return username, password

def handle_login(driver):
    """Handle the login process with browser interaction."""
    try:
        print("\nWaiting for login page to load...")
        print("\nPlease follow these steps:")
        print("1. In the browser window:")
        print("   - Enter your credentials for", driver.current_url)
        print("   - Click 'Sign In' or 'Log In'")
        print("\n2. If you have multiple profiles:")
        print("   - Select your Admin profile from the profile selection screen")
        print("   - Wait for the profile to load")
        print("\n3. Once you see the dashboard screen:")
        print("   - Ensure all elements have loaded")
        print("   - Return to this window")
        print("   - Press Enter to continue")
        
        # Wait for initial page load
        wait = WebDriverWait(driver, 20)
        
        # Let user handle login manually
        input("\nPress Enter when you're on the dashboard screen...")
        
        # Verify we're logged in (not on login page)
        try:
            current_url = driver.current_url
            if 'login' in current_url.lower():
                print("\nError: Still on login page. Please make sure you're logged in.")
                input("Press Enter to exit...")
                return False
            else:
                print("Login verified!")
                return True
        except Exception as e:
            print("\nError verifying login status:")
            print(str(e))
            print("Current URL:", driver.current_url)
            input("Press Enter to exit...")
            return False
        
    except Exception as e:
        print("\nError during login process:")
        print(str(e))
        print("\nCurrent URL:", driver.current_url)
        print("\nDebug info:")
        print(f"Page title: {driver.title}")
        input("Press Enter to exit...")
        return False

def scrape_fluxx_data(driver, url):
    try:
        # Clean up the URL
        url = url.strip()
        if url.startswith('@'):
            url = url[1:]
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
            
        print(f"\nNavigating to: {url}")
        
        # Try multiple navigation methods
        try:
            print("Attempting direct navigation...")
            driver.get(url)
            if driver.current_url == "data:,":
                raise Exception("Direct navigation failed")
        except:
            print("Trying JavaScript navigation...")
            driver.execute_script(f"window.location.href = '{url}';")
            time.sleep(2)
            
        current_url = driver.current_url
        print(f"Current URL: {current_url}")
        
        if current_url == "data:,":
            raise Exception("Unable to navigate to the URL")
            
        # Handle login
        if not handle_login(driver):
            print("Login failed. Exiting...")
            return None
            
        # TODO: Add actual scraping logic here
        pass
        
    except Exception as e:
        print(f"\nError during navigation/scraping:")
        print(str(e))
        print(f"Attempted URL: {url}")
        print(f"Current URL: {driver.current_url if driver else 'Not available'}")
        input("\nPress Enter to continue...")
        return None

def print_divider():
    """Print a visual divider line"""
    print("\n" + "=" * 50 + "\n")

def wait_with_spinner(message, action_func, *args, **kwargs):
    """Execute a function while showing a loading spinner"""
    stop_spinner = threading.Event()
    spinner_thread = threading.Thread(
        target=show_spinner, 
        args=(stop_spinner, message)
    )
    spinner_thread.start()
    
    try:
        result = action_func(*args, **kwargs)
        stop_spinner.set()
        spinner_thread.join()
        return result
    except Exception as e:
        stop_spinner.set()
        spinner_thread.join()
        raise e

def wait_for_dashboard(driver, timeout=60):
    """Wait for dashboard to load and verify we're logged in"""
    try:
        print("\n" + "=" * 80)
        print("\nLogin Instructions:")
        print("\n1. Enter your credentials in the browser window")
        print("2. Click 'Sign In' or 'Log In'")
        print("3. If prompted, select your Admin profile")
        print("\nIMPORTANT:")
        print("- Wait for the dashboard to fully load")
        print("- Once logged in, avoid clicking any elements in the browser")
        print("- The script will automatically navigate to required sections")
        print("\n" + "=" * 80)
        
        def wait_for_admin():
            wait = WebDriverWait(driver, timeout)
            return wait.until(
                EC.presence_of_element_located((By.CSS_SELECTOR, 'a.to-admin-panel[href="/?db=config"]'))
            )
            
        wait_with_spinner("Waiting for login...", wait_for_admin)
        print("\nLogin successful!")
        return True
        
    except TimeoutException:
        print("\nLogin timed out")
        return False
    except Exception as e:
        print(f"\nLogin error: {str(e)}")
        return False

def navigate_to_forms(driver):
    """Navigate to Forms dashboard"""
    try:
        def nav_action():
            wait = WebDriverWait(driver, 10)
            forms_link = wait.until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, 'li.item a.to-dashboard[href*="/client_stores/"]'))
            )
            # Verify it's the Forms link by checking text content
            if forms_link.text.strip() == 'Forms':
                forms_link.click()
                time.sleep(2)  # Wait for dashboard to load
            else:
                # If first link isn't Forms, find all dashboard links and click the Forms one
                dashboard_links = driver.find_elements(By.CSS_SELECTOR, 'li.item a.to-dashboard[href*="/client_stores/"]')
                for link in dashboard_links:
                    if link.text.strip() == 'Forms':
                        link.click()
                        time.sleep(2)
                        break
            
        wait_with_spinner("Navigating to Forms dashboard...", nav_action)
        print("Successfully navigated to Forms dashboard")
        return True
        
    except Exception as e:
        print(f"Error navigating to Forms dashboard: {str(e)}")
        return False

def navigate_to_admin(driver):
    """Navigate to Admin Panel"""
    try:
        def nav_action():
            wait = WebDriverWait(driver, 10)
            admin_button = wait.until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, 'a.to-admin-panel[href="/?db=config"]'))
            )
            print("Found Admin Panel button")
            admin_button.click()
            wait.until(EC.url_contains('db=config'))
            
        wait_with_spinner("Navigating to Admin Panel...", nav_action)
        print("Successfully navigated to Admin Panel")
        
        # After reaching Admin Panel, ensure we're on Forms dashboard
        return navigate_to_forms(driver)
        
    except Exception as e:
        print(f"Error navigating to Admin Panel: {str(e)}")
        return False

def navigate_to_workflows(driver):
    """Navigate to Workflow dashboard"""
    try:
        def nav_action():
            wait = WebDriverWait(driver, 10)
            workflow_link = wait.until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, 'li.item a.to-dashboard[href*="/client_stores/"]'))
            )
            # Verify it's the Workflow link by checking text content
            if workflow_link.text.strip() == 'Workflow':
                workflow_link.click()
                time.sleep(2)  # Wait for dashboard to load
            else:
                # If first link isn't Workflow, find all dashboard links and click the Workflow one
                dashboard_links = driver.find_elements(By.CSS_SELECTOR, 'li.item a.to-dashboard[href*="/client_stores/"]')
                for link in dashboard_links:
                    if link.text.strip() == 'Workflow':
                        link.click()
                        time.sleep(2)
                        break
            
        wait_with_spinner("Navigating to Workflow dashboard...", nav_action)
        print("Successfully navigated to Workflow dashboard")
        return True
        
    except Exception as e:
        print(f"Error navigating to Workflow dashboard: {str(e)}")
        return False

def scan_model_workflows(driver, models_data):
    """Scan workflow states and actions for each model"""
    try:
        print("\n" + "=" * 80)
        print("\n                     Scanning Model Workflows")
        print("\n" + "=" * 80)
        print("\nThis process will scan workflows from all models automatically.")
        print("This may take several minutes to complete depending on the number of models.")
        print("\nIMPORTANT:")
        print("- You can stop the process at any time by pressing Ctrl+C")
        print("- Please do not interact with the browser while the process is running")
        print("- The browser will automatically handle all interactions")
        print("\n" + "-" * 80)
        verify = input("\nWould you like to proceed with scanning workflows? (y/n): ").strip().lower()
        
        if verify != 'y':
            print("\nSkipping workflow scanning process.")
            return models_data
            
        # First navigate to Workflow dashboard
        if not navigate_to_workflows(driver):
            print("\nError: Could not navigate to Workflow dashboard")
            return models_data
            
        # Initialize counters for progress bar
        total_models = len(models_data)
        current_model = 0
        
        # Clear screen and show initial progress
        os.system('cls' if os.name == 'nt' else 'clear')
        print("\n" + "=" * 80)
        print("\n                     Scanning Model Workflows")
        print("\nPress Ctrl+C to stop the process at any time")
        print("Please wait while workflows are scanned from all models...")
        print("\n" + "=" * 80 + "\n")
        
        # Process each model
        for model_name, model_data in models_data.items():
            try:
                current_model += 1
                progress = (current_model / total_models) * 100
                
                # Update progress bar
                bar_length = 50
                filled_length = int(bar_length * current_model // total_models)
                bar = '=' * filled_length + '-' * (bar_length - filled_length)
                
                # Truncate model name if too long (limit to 20 chars)
                truncated_name = model_name[:20] + '...' if len(model_name) > 20 else model_name
                
                sys.stdout.write('\r' + ' ' * 100)  # Clear line
                sys.stdout.write('\rScanning Workflows: [{0}] {1:.1f}% ({2}/{3}) - {4}'.format(
                    bar, progress, current_model, total_models, truncated_name
                ))
                sys.stdout.flush()

                # Try different model name formats for the selector
                possible_ids = [
                    model_name,  # Original name
                    model_name.replace(" ", ""),  # No spaces
                    ''.join(word.capitalize() for word in model_name.split()),  # CamelCase
                    f"MacModelTypeDyn{model_name.replace(' ', '')}",  # Dynamic model format
                    model_name.replace(" ", "_")  # Underscores
                ]
                
                model_element = None
                for model_id in possible_ids:
                    selector = f"div.link.is-admin[data-id='{model_id}']"
                    elements = driver.find_elements(By.CSS_SELECTOR, selector)
                    if elements:
                        model_element = elements[0]
                        break
                
                if model_element:
                    # Click the model
                    driver.execute_script("arguments[0].click();", model_element)
                    time.sleep(2)  # Wait for potential UI update
                    
                    # Get themes from the model data
                    themes = model_data.get('themes', {})
                    if not themes:
                        models_data[model_name]['workflow'] = {'themes': {}}
                        continue
                    
                    workflow_data = {'themes': {}}
                    
                    # Find all theme links excluding "New Theme" and "Retired Themes"
                    theme_links = driver.find_elements(By.CSS_SELECTOR, 
                        "li.icon:not(.new-theme):not(.retired-themes) > a.link[title]")
                    
                    for theme_link in theme_links:
                        try:
                            theme_name = theme_link.get_attribute('title')
                            if not theme_name or theme_name not in themes:
                                continue
                                
                            # Click the theme
                            driver.execute_script("arguments[0].click();", theme_link)
                            time.sleep(2)  # Wait for theme to load
                            
                            try:
                                # Wait for workflow container
                                wait = WebDriverWait(driver, 10)
                                workflow_container = wait.until(
                                    EC.presence_of_element_located((By.CSS_SELECTOR, "div.listing[data-type='listing'][data-src='/machine_states']"))
                                )
                                
                                # Find all states in the workflow container
                                states = workflow_container.find_elements(By.CSS_SELECTOR, "li.entry[data-model-id]")
                                
                                if not states:
                                    workflow_data['themes'][theme_name] = {
                                        'workflow_id': None,
                                        'states': []
                                    }
                                    continue
                                
                                # Get workflow ID from any new event link
                                workflow_id = None
                                new_event_links = workflow_container.find_elements(By.CSS_SELECTOR, "a.new-event")
                                if new_event_links:
                                    href = new_event_links[0].get_attribute('href')
                                    match = re.search(r'machine_workflow_id=(\d+)', href)
                                    if match:
                                        workflow_id = match.group(1)
                                
                                # Process each state
                                theme_states = []
                                for state in states:
                                    try:
                                        # Get state header with both display and internal names
                                        state_header = state.find_element(By.CSS_SELECTOR, "h2").text.strip()
                                        
                                        # Parse display and internal names
                                        match = re.match(r'(.*?)\s*\((.*?)\)', state_header)
                                        if match:
                                            display_name, internal_name = match.groups()
                                        else:
                                            display_name = internal_name = state_header
                                            
                                        # Click state to get validation blocks
                                        state_link = state.find_element(By.CSS_SELECTOR, "a.to-detail")
                                        driver.execute_script("arguments[0].click();", state_link)
                                        time.sleep(1)
                                        
                                        # Wait for state details
                                        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "form.machine_state")))
                                        
                                        # Get all validation blocks
                                        validation_blocks = {}
                                        try:
                                            current_before = driver.find_element(By.CSS_SELECTOR, 
                                                "#machine_state_unsafe_before_validation_enter").get_attribute("value")
                                            if current_before and current_before.strip():
                                                validation_blocks['current_before_validation'] = current_before.strip()
                                                
                                            draft_before = driver.find_element(By.CSS_SELECTOR,
                                                "#machine_state_draft_before_validation_enter").get_attribute("value")
                                            if draft_before and draft_before.strip():
                                                validation_blocks['draft_before_validation'] = draft_before.strip()
                                                
                                            current_after = driver.find_element(By.CSS_SELECTOR,
                                                "#machine_state_unsafe_after_enter").get_attribute("value")
                                            if current_after and current_after.strip():
                                                validation_blocks['current_after_enter'] = current_after.strip()
                                                
                                            draft_after = driver.find_element(By.CSS_SELECTOR,
                                                "#machine_state_draft_after_enter").get_attribute("value")
                                            if draft_after and draft_after.strip():
                                                validation_blocks['draft_after_enter'] = draft_after.strip()
                                        except Exception as e:
                                            print(f"Error getting validation blocks: {str(e)}")
                                        
                                        # Get actions for this state
                                        actions = []
                                        action_elements = state.find_elements(By.CSS_SELECTOR, 
                                            "ul.events > li:not(:last-child) > a.to-detail")
                                        
                                        for action in action_elements:
                                            try:
                                                action_name = action.text.strip()
                                                if action_name and not action_name == '+':
                                                    # Click action to get details
                                                    driver.execute_script("arguments[0].click();", action)
                                                    time.sleep(1)
                                                    
                                                    # Wait for action details
                                                    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "form.machine_event")))
                                                    
                                                    # Get to state
                                                    to_state = None
                                                    try:
                                                        to_state_select = driver.find_element(By.CSS_SELECTOR, "#machine_event_to_state_id")
                                                        to_state = to_state_select.find_element(By.CSS_SELECTOR, "option[selected]").text.strip()
                                                    except:
                                                        pass
                                                    
                                                    # Get guard instructions
                                                    guard_instructions = None
                                                    try:
                                                        guard = driver.find_element(By.CSS_SELECTOR,
                                                            "#machine_event_unsafe_guard").get_attribute("value")
                                                        if guard and guard.strip():
                                                            guard_instructions = guard.strip()
                                                    except:
                                                        pass
                                                    
                                                    # Get draft guard instructions
                                                    draft_guard = None
                                                    try:
                                                        draft = driver.find_element(By.CSS_SELECTOR,
                                                            "#machine_event_draft_guard").get_attribute("value")
                                                        if draft and draft.strip():
                                                            draft_guard = draft.strip()
                                                    except:
                                                        pass
                                                    
                                                    action_data = {
                                                        'name': action_name,
                                                        'to_state': to_state,
                                                        'guard_instructions': guard_instructions,
                                                        'draft_guard': draft_guard
                                                    }
                                                    actions.append(action_data)
                                            except Exception as e:
                                                print(f"Error processing action {action_name}: {str(e)}")
                                                continue
                                        
                                        state_data = {
                                            'display_name': display_name.strip(),
                                            'internal_name': internal_name.strip(),
                                            'validation_blocks': validation_blocks,
                                            'actions': actions
                                        }
                                        theme_states.append(state_data)
                                            
                                    except Exception as e:
                                        print(f"\nError processing state: {str(e)}")
                                        continue
                                
                                workflow_data['themes'][theme_name] = {
                                    'workflow_id': workflow_id,
                                    'states': theme_states
                                }

                            except TimeoutException:
                                print(f"Workflow container not found for theme: {theme_name}")
                                workflow_data['themes'][theme_name] = {
                                    'workflow_id': None,
                                    'states': []
                                }
                                continue
                                
                        except Exception as e:
                            print(f"\nError processing theme {theme_name}: {str(e)}")
                            continue
                    
                    # Store workflow data in model dictionary
                    models_data[model_name]['workflow'] = workflow_data
                else:
                    # Initialize empty workflow data structure for models without workflows
                    models_data[model_name]['workflow'] = {'themes': {}}
                    
            except Exception as e:
                print(f"\nError processing model {model_name}: {str(e)}")
                models_data[model_name]['workflow'] = {'themes': {}}
                continue
                
        # Show completion
        sys.stdout.write('\r' + ' ' * 100)  # Clear line
        sys.stdout.write('\rWorkflow scanning complete!')
        sys.stdout.flush()
        print("\n" + "=" * 80)
        
        return models_data
        
    except Exception as e:
        print(f"\nError during workflow scanning: {str(e)}")
        return models_data

def navigate_to_card_settings(driver):
    """Navigate to Card Settings section"""
    try:
        def nav_action():
            wait = WebDriverWait(driver, 10)
            
            # First find and click the dashboard picker combo if needed
            try:
                combo = driver.find_element(By.CSS_SELECTOR, "li.combo")
                driver.execute_script("arguments[0].click();", combo)
                time.sleep(1)  # Wait for dropdown to open
            except:
                pass
            
            # Find and click the Card Settings link
            card_settings_link = wait.until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, 'li.item a.to-dashboard[href*="/client_stores/"][data-updated="NaN/NaN/NaN"]'))
            )
            
            # Make sure it's the Card Settings link
            links = driver.find_elements(By.CSS_SELECTOR, 'li.item a.to-dashboard[href*="/client_stores/"]')
            for link in links:
                if link.text.strip() == "Card Settings":
                    driver.execute_script("arguments[0].click();", link)
                    time.sleep(2)  # Wait for page to load
                    return
                    
            raise Exception("Could not find Card Settings link")
            
        wait_with_spinner("Navigating to Card Settings...", nav_action)
        print("Successfully navigated to Card Settings")
        return True
        
    except Exception as e:
        print(f"Error navigating to Card Settings: {str(e)}")
        return False

def scan_methods(driver, models_data):
    """Scan methods from all models"""
    try:
        print("\n" + "=" * 80)
        print("\n                     Scanning Model Methods")
        print("\n" + "=" * 80)
        print("\nThis process will scan methods from all models automatically.")
        print("This may take several minutes to complete depending on the number of models.")
        print("\nIMPORTANT:")
        print("- You can stop the process at any time by pressing Ctrl+C")
        print("- Please do not interact with the browser while the process is running")
        print("- The browser will automatically handle all interactions")
        print("\n" + "-" * 80)
        verify = input("\nWould you like to proceed with scanning methods? (y/n): ").strip().lower()
        
        if verify != 'y':
            print("\nSkipping method scanning process.")
            return models_data
            
        # Initialize WebDriverWait
        wait = WebDriverWait(driver, 10)

        # First navigate to Card Settings
        if not navigate_to_card_settings(driver):
            print("\nError: Could not navigate to Card Settings")
            return models_data
            
        # Initialize counters for progress bar
        total_models = len(models_data)
        current_model = 0
        
        # Clear screen and show initial progress
        os.system('cls' if os.name == 'nt' else 'clear')
        print("\n" + "=" * 80)
        print("\n                     Scanning Model Methods")
        print("\nPress Ctrl+C to stop the process at any time")
        print("Please wait while methods are scanned from all models...")
        print("\n" + "=" * 80 + "\n")
        
        # Process each model
        for model_name, model_data in models_data.items():
            try:
                current_model += 1
                progress = (current_model / total_models) * 100
                
                # Update progress bar with model name
                bar_length = 50
                filled_length = int(bar_length * current_model // total_models)
                bar = '=' * filled_length + '-' * (bar_length - filled_length)
                
                status_text = f"Scanning Methods: [{bar}] {progress:.1f}% ({current_model}/{total_models}) - {model_name}"
                
                # Clear line and write status
                sys.stdout.write('\r' + ' ' * 100)  # Clear line
                sys.stdout.write('\r' + status_text)
                sys.stdout.flush()

                # Try different model name formats for the selector
                possible_ids = [
                    model_name,  # Original name
                    model_name.replace(" ", ""),  # No spaces
                    ''.join(word.capitalize() for word in model_name.split()),  # CamelCase
                    f"MacModelTypeDyn{model_name.replace(' ', '')}",  # Dynamic model format
                    model_name.replace(" ", "_")  # Underscores
                ]
                
                model_element = None
                for model_id in possible_ids:
                    selector = f"div.link.is-admin[data-id='{model_id}']"
                    elements = driver.find_elements(By.CSS_SELECTOR, selector)
                    if elements:
                        model_element = elements[0]
                        break
                
                if model_element:
                    # Click the model
                    driver.execute_script("arguments[0].click();", model_element)
                    time.sleep(2)  # Wait for potential UI update
                    
                    # First check if Methods tab exists
                    methods_tab = None
                    try:
                        # Try finding the Methods tab by text content first
                        tabs = driver.find_elements(By.CSS_SELECTOR, "ul.dock-tabs li a.ui-tabs-anchor")
                        for tab in tabs:
                            if tab.text.strip() == "Methods":
                                methods_tab = tab
                                break
                                
                        if not methods_tab:
                            # Try alternate selector
                            methods_tab = driver.find_element(By.CSS_SELECTOR, "ul.dock-tabs li a[href*='fluxx-card'][id*='ui-id']")
                            
                    except:
                        pass
                            
                    if not methods_tab:
                        models_data[model_name]['methods'] = []
                        continue

                    # Click the Methods tab
                    driver.execute_script("arguments[0].click();", methods_tab)
                    time.sleep(1)  # Wait for tab content to load
                    
                    # Wait for methods container with updated selector
                    try:
                        methods_container = wait.until(
                            EC.presence_of_element_located((By.CSS_SELECTOR, "div.listing.area[data-type='listing'][data-src='/model_methods']"))
                        )
                        
                        # Find all method entries in the list
                        method_entries = methods_container.find_elements(By.CSS_SELECTOR, "ul.list > li.entry")
                        
                        if not method_entries:
                            models_data[model_name]['methods'] = []
                            continue
                        
                        model_methods = []
                        for method_entry in method_entries:
                            try:
                                # Get method name from the h2 in the list entry
                                method_name = method_entry.find_element(By.CSS_SELECTOR, "h2").text.strip()
                                
                                # Find and click the method link to open details
                                method_link = method_entry.find_element(By.CSS_SELECTOR, "a.to-detail")
                                driver.execute_script("arguments[0].click();", method_link)
                                time.sleep(1)  # Wait for details to load
                                
                                # Wait for detail area to be visible and loaded
                                detail_area = wait.until(
                                    EC.presence_of_element_located((By.CSS_SELECTOR, "div.detail.area[data-type='detail']"))
                                )
                                
                                # Get method type from select dropdown
                                method_type = ""
                                try:
                                    type_select = detail_area.find_element(By.CSS_SELECTOR, "#model_method_method_type")
                                    selected_option = type_select.find_element(By.CSS_SELECTOR, "option[selected]")
                                    method_type = selected_option.text.strip()
                                except:
                                    pass
                                
                                # Get current dynamic method code
                                current_code = ""
                                try:
                                    code_element = detail_area.find_element(By.CSS_SELECTOR, "#model_method_unsafe_dyn_method")
                                    current_code = code_element.get_attribute("value").strip()
                                except:
                                    pass
                                
                                # Get draft dynamic method code
                                draft_code = ""
                                try:
                                    draft_element = detail_area.find_element(By.CSS_SELECTOR, "#model_method_draft_dyn_method")
                                    draft_code = draft_element.get_attribute("value").strip()
                                except:
                                    pass
                                
                                method_data = {
                                    'name': method_name,
                                    'type': method_type,
                                    'current_code': current_code,
                                    'draft_code': draft_code
                                }
                                model_methods.append(method_data)
                                    
                            except Exception:
                                continue
                        
                        # Store methods directly in model data
                        models_data[model_name]['methods'] = model_methods
                        
                    except TimeoutException:
                        models_data[model_name]['methods'] = []
                        continue
                            
                else:
                    # Initialize empty methods list for models without methods
                    models_data[model_name]['methods'] = []
                    
            except Exception as e:
                print(f"\nError processing model {model_name}: {str(e)}")
                models_data[model_name]['methods'] = []
                continue
        
        # Show completion
        sys.stdout.write('\r' + ' ' * 100)
        sys.stdout.write('\rMethod scanning complete!')
        sys.stdout.flush()
        print("\n" + "=" * 80)
        
        return models_data
        
    except Exception as e:
        print(f"\nError during method scanning: {str(e)}")
        return models_data

# HTML Structure Reference for Forms Section:
#
# Models:
# - Located in: #iconList > ul
# - Model name from: ul[id] attribute (replace _ with space, title case)
# Example: <ul id="custom_ui_enhancements" class="toggle-class open" data-click-when-opened=".scroll-to-card">
#
# Themes:
# - Located in: ul > li.icon[data-card-uid]
# - Theme name from: li.icon > a.link.scroll-to-card > span.label[role='menuitem']
# Example:
# <li class="icon" data-card-uid="19980">
#   <a class="link scroll-to-card" href="#fluxx-card-26">
#     <span class="label" role="menuitem">Animation Examples</span>
#
# Views:
# - Container: li.icon > div.listing[data-type='listing'][data-src='/stencils']
# - View items: div.listing > ul.list > li.entry:not(.non-entry)
# - View name from: li.entry > a.to-detail > div.label
# Example:
# <div class="listing" data-type="listing" data-src="/stencils" ...>
#   <ul class="list">
#     <li class="active entry selected" data-model-id="35727">
#       <a class="to-detail" href="/stencils/35727">
#         <div class="label">Gallery</div>

def wait_for_forms_and_parse(driver, max_retries=3):
    """Wait for Forms section to load and parse content with retry logic"""
    try:
        # Clear screen and show header
        os.system('cls' if os.name == 'nt' else 'clear')
        print("\n" + "=" * 80)
        print("\n                     Scanning Models and Themes")
        print("\n" + "=" * 80)
        
        wait = WebDriverWait(driver, 10)
        
        # First verify Forms section exists
        wait.until(EC.presence_of_element_located(
            (By.CSS_SELECTOR, "#iconList")
        ))
        
        # Wait for at least one model to be loaded
        try:
            wait.until(lambda d: len(d.find_elements(By.CSS_SELECTOR, "#iconList > ul[id]")) > 0)
        except TimeoutException:
            wait = WebDriverWait(driver, 30)
            wait.until(lambda d: len(d.find_elements(By.CSS_SELECTOR, "#iconList > ul[id]")) > 0)
            time.sleep(2)
        
        while True:
            model_list = driver.find_elements(By.CSS_SELECTOR, "#iconList > ul[id]")
            model_list = [model for model in model_list if model.get_attribute("id")]
            model_count = len(model_list)
            
            # Get the height of the iconList container
            icon_list = driver.find_element(By.CSS_SELECTOR, "#iconList")
            list_height = driver.execute_script("return arguments[0].scrollHeight", icon_list)
            
            # Scroll through the list to ensure all models are loaded
            current_scroll = 0
            scroll_step = 500
            
            while current_scroll < list_height:
                driver.execute_script(f"arguments[0].scrollTop = {current_scroll}", icon_list)
                time.sleep(0.5)
                current_scroll += scroll_step
            
            driver.execute_script("arguments[0].scrollTop = 0", icon_list)
            
            model_list = driver.find_elements(By.CSS_SELECTOR, "#iconList > ul[id]")
            model_list = [model for model in model_list if model.get_attribute("id")]
            new_count = len(model_list)
            
            if new_count > model_count:
                model_count = new_count
                continue
            
            print(f"\nFound {model_count} models.")
            verify = input("Does this count appear correct? (y/n): ").strip().lower()
            
            if verify == 'y':
                break
            print("\nWaiting for more models to load...")
            time.sleep(3)
                
        # Initialize dictionary to store model data
        models = {}
        current_model = 0
        total_models = len(model_list)
        
        # Print initial progress bar
        sys.stdout.write("\rScanning Models: [--------------------------------------------------] 0.0% (0/{})".format(total_models))
        sys.stdout.flush()
        
        # Process each model
        for model_ul in model_list:
            try:
                current_model += 1
                progress = (current_model / total_models) * 100
                
                # Create progress bar
                bar_length = 50
                filled_length = int(bar_length * current_model // total_models)
                bar = '=' * filled_length + '-' * (bar_length - filled_length)
                
                # Update progress bar
                sys.stdout.write('\r' + ' ' * 100)
                sys.stdout.write('\rScanning Models: [{0}] {1:.1f}% ({2}/{3})'.format(
                    bar, progress, current_model, total_models
                ))
                sys.stdout.flush()
                
                # Get model name from the UL id attribute
                model_id = model_ul.get_attribute("id")
                if not model_id:
                    continue
                    
                model_name = model_id.replace('_', ' ').title()
                
                # Find the "New Theme" link to extract model type
                model_type = None
                try:
                    # Try multiple selectors to find model type
                    selectors = [
                        "a.link.to-modal[href*='model_theme[model_type]']",
                        "a.link[href*='model_theme[model_type]']",
                        "a[href*='model_theme[model_type]']"
                    ]
                    
                    for selector in selectors:
                        try:
                            new_theme_link = model_ul.find_element(By.CSS_SELECTOR, selector)
                            if new_theme_link:
                                href = new_theme_link.get_attribute("href")
                                # Extract model type from URL parameter
                                match = re.search(r'model_theme\[model_type\]=(\w+)', href)
                                if match:
                                    model_type = match.group(1)
                        except:
                            continue
                except:
                    pass
                
                # Check if model is dynamic
                is_dynamic = model_type and model_type.startswith('MacModelTypeDyn')
                
                models[model_name] = {
                    'type': model_type,
                    'is_dynamic': is_dynamic,
                    'themes': {}
                }
                
                # Find themes within this model's UL
                theme_items = model_ul.find_elements(By.CSS_SELECTOR, 
                    "li.icon[data-card-uid]")
                
                for theme in theme_items:
                    try:
                        theme_label = theme.find_element(By.CSS_SELECTOR, 
                            "a.link.scroll-to-card span.label").text
                        theme_name = theme_label.get_attribute("textContent").strip()
                        
                        if theme_name and theme_name not in ['New Theme', 'Retired Themes', 'Export', 'Filter', 'Visualizations']:
                            models[model_name]['themes'][theme_name] = {'views': []}
                            
                            # Find and process views
                            listing_div = theme.find_element(By.CSS_SELECTOR, 
                                "div.listing[data-type='listing'][data-src='/stencils']")
                            
                            views = listing_div.find_elements(By.CSS_SELECTOR,
                                "ul.list > li.entry:not(.non-entry)")
                            
                            for view in views:
                                try:
                                    label_div = view.find_element(By.CSS_SELECTOR, 
                                        "a.to-detail > div.label")
                                    view_name = label_div.get_attribute("textContent").strip()
                                    
                                    if view_name and view_name != 'New View':
                                        models[model_name]['themes'][theme_name]['views'].append(view_name)
                                except:
                                    continue
                    except:
                        continue
                        
            except:
                continue
        
        # Show completion
        sys.stdout.write('\r' + ' ' * 100)
        sys.stdout.write('\rModel scanning complete!')
        sys.stdout.flush()
        print("\n" + "=" * 80)
        
        return models
        
    except Exception as e:
        print(f"\nError parsing Forms section: {str(e)}")
        return None

def generate_word_document(models_data, site_url=None):
    """Generate a Word document using the Social Edge template format"""
    try:
        # Create document
        doc = Document()
        
        # Set standard margins (1 inch = 1440 twips)
        sections = doc.sections
        for section in sections:
            section.left_margin = int(1.25 * 1440)  # 1.25 inch left margin
            section.right_margin = int(1.25 * 1440)  # 1.25 inch right margin
            section.top_margin = int(1.25 * 1440)  # 1.25 inch top margin
            section.bottom_margin = int(1.25 * 1440)  # 1.25 inch bottom margin
            # Set page dimensions
            section.page_width = Pt(8.5 * 72)  # 8.5 inches
            section.page_height = Pt(11 * 72)  # 11 inches
        
        # Set default font and paragraph styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_before = Pt(12)  # Space before paragraphs
        style.paragraph_format.space_after = Pt(12)   # Space after paragraphs
        style.paragraph_format.line_spacing = 1.15    # Line spacing
        
        # Update heading styles
        heading1 = doc.styles['Heading 1']
        heading1.font.name = 'Calibri'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.paragraph_format.space_before = Pt(24)  # Extra space before H1
        heading1.paragraph_format.space_after = Pt(12)
        
        heading2 = doc.styles['Heading 2']
        heading2.font.name = 'Calibri'
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.paragraph_format.space_before = Pt(18)  # Space before H2
        heading2.paragraph_format.space_after = Pt(12)
        
        # Add title with proper spacing
        title = doc.add_heading('Fluxx Build Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(36)  # Extra space at top
        title.paragraph_format.space_after = Pt(24)
        
        # Add subtitle with URL if provided
        if site_url:
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.paragraph_format.space_before = Pt(12)
            subtitle.paragraph_format.space_after = Pt(24)
            subtitle_text = subtitle.add_run(f"Generated for: {site_url}")
            subtitle_text.font.size = Pt(12)
            subtitle_text.font.name = 'Calibri'
            
        # Add generation timestamp
        timestamp = doc.add_paragraph()
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp.paragraph_format.space_before = Pt(12)
        timestamp.paragraph_format.space_after = Pt(36)  # Extra space after header
        timestamp_text = timestamp.add_run(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp_text.font.size = Pt(10)
        timestamp_text.font.name = 'Calibri'
        
        doc.add_page_break()
        
        # Process each model
        for model_name, model_data in models_data.items():
            try:
                # Add spacing before model section
                doc.add_paragraph().paragraph_format.space_before = Pt(24)
                
                # Create table with specific width and center alignment
                table = doc.add_table(rows=9, cols=2)
                table.style = 'Table Grid'
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Set table width to 90% of page width
                section = doc.sections[0]
                available_width = section.page_width - (section.left_margin + section.right_margin)
                table_width = int(available_width * 0.9)
                table.width = table_width
                
                # Set column widths (30% for labels, 70% for content)
                for i, width in enumerate([0.3, 0.7]):
                    for cell in table.columns[i].cells:
                        cell.width = int(table_width * width)
                
                # Define all row titles with their indices
                row_titles = {
                    0: 'Model Name',  # Header row
                    1: 'Themes:',
                    2: 'Workflow:',
                    4: 'Method:',
                    5: 'Before New / After Create:',
                    6: 'Before Validation / After Enter / Guard Instructions:',
                    7: 'Documents:',
                    8: 'Embedded Cards / Dynamic Relationships:'
                }
                
                # Set up all row titles and initialize empty cells
                for row_idx, title in row_titles.items():
                    if row_idx > 0:  # Skip header row
                        cell = table.rows[row_idx].cells[0]
                        cell.text = title
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].runs[0].font.name = 'Calibri'
                        # Initialize empty content cell
                        table.rows[row_idx].cells[1].text = ''
                
                # Model header (row 0)
                header_row = table.rows[0]
                header_row.cells[0].merge(header_row.cells[1])
                model_type = model_data.get('type', '')
                is_dynamic = model_data.get('is_dynamic', False)
                header_text = model_name
                if model_type:
                    header_text += f" ({model_type})"
                if is_dynamic:
                    header_text += " - Dynamic Model"
                header_cell = header_row.cells[0]
                header_para = header_cell.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_run = header_para.add_run(header_text)
                header_run.font.bold = True
                header_run.font.size = Pt(14)
                header_run.font.name = 'Calibri'
                
                # Process themes (row 1)
                themes = model_data.get('themes', {})
                theme_text = [f"{len(themes)} Themes Built"]
                for theme_name, theme_data in themes.items():
                    theme_line = [f"\n{theme_name}"]
                    views = theme_data.get('views', [])
                    for view in views:
                        theme_line.append(f"- {view}")
                    theme_text.append('\n'.join(theme_line))
                table.rows[1].cells[1].text = '\n'.join(theme_text)
                
                # Process workflows (row 2)
                workflow_data = model_data.get('workflow', {})
                if isinstance(workflow_data, dict):
                    themes = workflow_data.get('themes', {})
                    if themes:
                        workflow_text = []
                        for theme_name, theme_data in themes.items():
                            if isinstance(theme_data, dict):
                                states = theme_data.get('states', [])
                                if states:
                                    workflow_text.append(f"\nTheme: {theme_name}")
                                    for state in states:
                                        if isinstance(state, dict):
                                            display_name = state.get('display_name', '')
                                            internal_name = state.get('internal_name', '')
                                            workflow_text.append(f"• {display_name} ({internal_name})")
                                            
                                            actions = state.get('actions', [])
                                            for action in actions:
                                                if isinstance(action, dict):
                                                    action_text = f"  - {action.get('name', '')}"
                                                    to_state = action.get('to_state')
                                                    if to_state:
                                                        action_text += f" [To State -> {to_state}]"
                                                    workflow_text.append(action_text)
                        
                        if workflow_text:
                            workflow_cell = table.rows[2].cells[1]
                            workflow_cell.text = ''
                            for line in workflow_text:
                                para = workflow_cell.add_paragraph()
                                if line.startswith('\nTheme:'):
                                    run = para.add_run(line.strip())
                                    run.bold = True
                                elif line.startswith('  -'):
                                    run = para.add_run(line)
                                    run.font.size = Pt(9)
                                    run.font.italic = True
                                    run.font.color.rgb = RGBColor(128, 128, 128)
                                else:
                                    run = para.add_run(line)
                                para.paragraph_format.space_after = Pt(0)
                                para.paragraph_format.space_before = Pt(0)
                    else:
                        table.rows[2].cells[1].text = "No workflow states configured"
                else:
                    table.rows[2].cells[1].text = "No workflow configuration"
                
                # Process Before New / After Create (row 5)
                before_after_cell = table.rows[5].cells[1]
                code_text = []
                
                # Get themes from model data
                themes = model_data.get('themes', {})
                if themes:
                    has_code = False
                    for theme_name, theme_data in themes.items():
                        if isinstance(theme_data, dict):
                            theme_code = theme_data.get('code', {})
                            if theme_code:
                                theme_has_code = False
                                # Order blocks consistently
                                block_order = [
                                    ('current_before_new', 'Current Before New Block'),
                                    ('draft_before_new', 'Draft Before New Block'),
                                    ('current_after_create', 'Current After Create Block'),
                                    ('draft_after_create', 'Draft After Create Block')
                                ]
                                
                                for block_key, block_label in block_order:
                                    if block_key in theme_code and theme_code[block_key] and theme_code[block_key] != "N/A":
                                        if not theme_has_code:
                                            code_text.append(f"\nTheme: {theme_name}")
                                            theme_has_code = True
                                            has_code = True
                                        code_text.append(f"\n{block_label}:")
                                        code_text.append(theme_code[block_key])
                    
                    if not has_code:
                        code_text = ["No Before New or After Create blocks configured"]
                else:
                    code_text = ["No themes configured"]
                
                # Write to cell with proper formatting
                before_after_cell.text = ''  # Clear existing text
                for line in code_text:
                    para = before_after_cell.add_paragraph()
                    if line.startswith('\nTheme:'):
                        run = para.add_run(line.strip())
                        run.bold = True
                    elif line.endswith('Block:'):
                        run = para.add_run(line)
                        run.italic = True
                    elif line.startswith('No '):  # Handle "No blocks configured" messages
                        run = para.add_run(line)
                        run.italic = True
                    else:
                        run = para.add_run(line)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(128, 128, 128)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.space_before = Pt(0)
                
                # Process validation blocks and guard instructions (row 6)
                validation_cell = table.rows[6].cells[1]
                if isinstance(workflow_data, dict):
                    themes = workflow_data.get('themes', {})
                    if themes:
                        validation_text = []
                        validation_text.append("VALIDATION BLOCKS:")
                        
                        # Process validation blocks
                        for theme_name, theme_data in themes.items():
                            if isinstance(theme_data, dict):
                                states = theme_data.get('states', [])
                                if states:
                                    has_validation = False
                                    for state in states:
                                        if isinstance(state, dict):
                                            validation_blocks = state.get('validation_blocks', {})
                                            if validation_blocks:
                                                if not has_validation:
                                                    validation_text.append(f"\nTheme: {theme_name}")
                                                    has_validation = True
                                                
                                                display_name = state.get('display_name', '')
                                                internal_name = state.get('internal_name', '')
                                                validation_text.append(f"\nState: {display_name} ({internal_name})")
                                                
                                                # Order validation blocks consistently
                                                block_order = [
                                                    ('current_before_validation', 'Current Before Validation'),
                                                    ('draft_before_validation', 'Draft Before Validation'),
                                                    ('current_after_enter', 'Current After Enter'),
                                                    ('draft_after_enter', 'Draft After Enter')
                                                ]
                                                
                                                for block_key, block_label in block_order:
                                                    if block_key in validation_blocks and validation_blocks[block_key]:
                                                        validation_text.append(f"\n{block_label}:")
                                                        validation_text.append(validation_blocks[block_key])
                                
                                # Process guard instructions
                                has_guards = False
                                guard_text = ["\nGUARD INSTRUCTIONS:"]
                                
                                for theme_name, theme_data in themes.items():
                                    if isinstance(theme_data, dict):
                                        states = theme_data.get('states', [])
                                        if states:
                                            theme_has_guards = False
                                            for state in states:
                                                if isinstance(state, dict):
                                                    actions = state.get('actions', [])
                                                    for action in actions:
                                                        if isinstance(action, dict):
                                                            guard = action.get('guard_instructions')
                                                            draft_guard = action.get('draft_guard')
                                                            
                                                            if guard or draft_guard:
                                                                if not theme_has_guards:
                                                                    guard_text.append(f"\nTheme: {theme_name}")
                                                                    theme_has_guards = True
                                                                    has_guards = True
                                                                
                                                                display_name = state.get('display_name', '')
                                                                internal_name = state.get('internal_name', '')
                                                                action_name = action.get('name', '')
                                                                
                                                                if guard:
                                                                    guard_text.append(f"\nGuard Instructions for {action_name} in {display_name} ({internal_name}):")
                                                                    guard_text.append(guard)
                                                                if draft_guard:
                                                                    guard_text.append(f"\nDraft Guard Instructions for {action_name} in {display_name} ({internal_name}):")
                                                                    guard_text.append(draft_guard)
                                
                                if has_guards:
                                    validation_text.extend(guard_text)
                                
                                if len(validation_text) > 1:
                                    validation_cell.text = ''
                                    for line in validation_text:
                                        para = validation_cell.add_paragraph()
                                        if line.endswith('BLOCKS:') or line.endswith('INSTRUCTIONS:'):
                                            run = para.add_run(line)
                                            run.bold = True
                                        elif line.startswith('\nTheme:'):
                                            run = para.add_run(line.strip())
                                            run.bold = True
                                        elif line.startswith('\nGuard Instructions for') or line.endswith('Block:'):
                                            run = para.add_run(line)
                                            run.italic = True
                                        elif line.startswith('\nState:'):
                                            run = para.add_run(line)
                                            run.bold = True
                                            run.italic = True
                                        else:
                                            run = para.add_run(line)
                                            run.font.name = 'Consolas'
                                            run.font.size = Pt(9)
                                            run.font.color.rgb = RGBColor(128, 128, 128)
                                        para.paragraph_format.space_after = Pt(0)
                                        para.paragraph_format.space_before = Pt(0)
                                else:
                                    validation_cell.text = "No validation blocks or guard instructions configured"
                            else:
                                validation_cell.text = "No workflow configuration"
                        else:
                            validation_cell.text = "No workflow configuration"
                    else:
                        validation_cell.text = "No workflow configuration"
                else:
                    validation_cell.text = "No workflow configuration"
                
                # Rows 3, 4, 7, and 8 remain empty for user input
                
                # Add spacing after table
                doc.add_paragraph().paragraph_format.space_after = Pt(24)
                
                # Add page break between models
                if model_name != list(models_data.keys())[-1]:
                    doc.add_page_break()
                    
                # Process methods (row 4)
                methods = model_data.get('methods', [])
                if methods:
                    methods_cell = table.rows[4].cells[1]
                    methods_cell.text = ''  # Clear existing text
                    
                    # Add header showing total methods count
                    header_para = methods_cell.add_paragraph()
                    header_run = header_para.add_run(f"{len(methods)} Methods Found")
                    header_run.bold = True
                    
                    # Process each method
                    for method in methods:
                        if isinstance(method, dict):
                            # Add method name and type
                            name_para = methods_cell.add_paragraph()
                            name_para.paragraph_format.space_before = Pt(12)
                            name_run = name_para.add_run(f"\n{method.get('name', '')}")
                            name_run.bold = True
                            if method.get('type'):
                                type_run = name_para.add_run(f" ({method.get('type')})")
                                type_run.italic = True
                            
                            # Add current code if exists
                            if method.get('current_code'):
                                code_para = methods_cell.add_paragraph()
                                code_para.paragraph_format.space_before = Pt(6)
                                code_run = code_para.add_run("Current Code:")
                                code_run.italic = True
                                code_para = methods_cell.add_paragraph()
                                code_run = code_para.add_run(method['current_code'])
                                code_run.font.name = 'Consolas'
                                code_run.font.size = Pt(9)
                                code_run.font.color.rgb = RGBColor(128, 128, 128)
                            
                            # Add draft code if exists
                            if method.get('draft_code'):
                                draft_para = methods_cell.add_paragraph()
                                draft_para.paragraph_format.space_before = Pt(6)
                                draft_run = draft_para.add_run("Draft Code:")
                                draft_run.italic = True
                                draft_para = methods_cell.add_paragraph()
                                draft_run = draft_para.add_run(method['draft_code'])
                                draft_run.font.name = 'Consolas'
                                draft_run.font.size = Pt(9)
                                draft_run.font.color.rgb = RGBColor(128, 128, 128)
                else:
                    table.rows[4].cells[1].text = "No methods found"
                
            except Exception as e:
                print(f"\nError processing model {model_name}: {str(e)}")
                continue
        
        # Save document
        timestamp_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f'fluxx_documentation_{timestamp_str}.docx'
        doc.save(filename)
        print(f"\nWord document saved as: {filename}")
        return filename
        
    except Exception as e:
        print(f"\nError generating Word document: {str(e)}")
        return None

def show_spinner(stop_event, message=""):
    """Show a simple spinner animation with a message"""
    spinner = ['|', '/', '-', '\\']  # Simple ASCII spinner
    i = 0
    while not stop_event.is_set():
        print(f"\r{spinner[i]} {message}", end='', flush=True)
        i = (i + 1) % len(spinner)
        time.sleep(0.1)
    print("\r", end='', flush=True)  # Clear the spinner line

def print_header():
    """Print the application header with clean formatting"""
    os.system('cls' if os.name == 'nt' else 'clear')
    
    # Social Edge logo ASCII art
    print("""
                                                                                           
                             
                    ;X.            
                  ;XX+             
                .XX+  ..:.         
              .XX+  xXXxxXXX
            .XXx  xXX.    .XX
            XX  ;XX. .XX+  XX
            Xx  .  .XXX  +XX.
            ;XXx.;XXx  ;XX;
              .+++:  :XX;
                   :XX+            
                  +X+            
                                                           
       Social Edge Consulting - 2025
   nick.reitan@socialedgeconsulting.com                    
                                                                   
    """)
    
    print("  Fluxx Build Documentation Automated Tool    ")
    print_divider()

def check_chrome_and_driver():
    """Diagnose Chrome and ChromeDriver versions and compatibility"""
    try:
        print("Please wait while your Google Chrome version is verified...")
        print("(This will take a few moments)")
        print_divider()
        
        # Create a threading event to control the spinner
        stop_spinner = threading.Event()
        spinner_thread = threading.Thread(
            target=show_spinner, 
            args=(stop_spinner, "Checking Chrome setup...")
        )
        spinner_thread.start()
        
        # Get Chrome version
        chrome_path = get_chrome_path()
        if not chrome_path:
            stop_spinner.set()
            spinner_thread.join()
            print("Error: Chrome not found. Please install Google Chrome.")
            return False
            
        chrome_version = get_chrome_version()
        
        # Update spinner message
        stop_spinner.set()
        spinner_thread.join()
        print(f"Chrome version detected: {chrome_version}")
        print_divider()
        
        # Check if chromedriver exists and is compatible
        driver_path = os.path.join(os.getcwd(), "chromedriver.exe")
        if os.path.exists(driver_path):
            stop_spinner = threading.Event()
            spinner_thread = threading.Thread(
                target=show_spinner, 
                args=(stop_spinner, "Verifying ChromeDriver compatibility...")
            )
            spinner_thread.start()
            
            try:
                service = Service(driver_path, log_path='NUL')
                options = webdriver.ChromeOptions()
                options.add_argument('--headless')
                options.add_argument('--log-level=3')  # Suppress console messages
                options.add_experimental_option('excludeSwitches', ['enable-logging'])
                driver = webdriver.Chrome(service=service, options=options)
                driver.quit()
                stop_spinner.set()
                spinner_thread.join()
                print("\nSetup complete!")
                print_divider()
                time.sleep(0.5)
                return True
            except Exception:
                stop_spinner.set()
                spinner_thread.join()
                print("\nUpdating ChromeDriver to match Chrome version...")
        else:
            print("\nSetting up ChromeDriver...")
            
        # Download matching ChromeDriver if needed
        stop_spinner = threading.Event()
        spinner_thread = threading.Thread(
            target=show_spinner, 
            args=(stop_spinner, "Downloading and installing ChromeDriver...")
        )
        spinner_thread.start()
        
        chrome_major = chrome_version.split('.')[0]
        driver_url = f"https://edgedl.me.gvt1.com/edgedl/chrome/chrome-for-testing/{chrome_version}/win64/chromedriver-win64.zip"
        
        try:
            response = requests.get(driver_url)
            response.raise_for_status()
            
            with open("chromedriver.zip", "wb") as f:
                f.write(response.content)
                
            with zipfile.ZipFile("chromedriver.zip", "r") as zip_ref:
                zip_ref.extractall()
                
            if os.path.exists("chromedriver.exe"):
                os.remove("chromedriver.exe")
            shutil.move("chromedriver-win64/chromedriver.exe", "chromedriver.exe")
            
            # Clean up
            os.remove("chromedriver.zip")
            shutil.rmtree("chromedriver-win64")
            
            stop_spinner.set()
            spinner_thread.join()
            print("Setup complete!")
            return True
            
        except Exception as e:
            stop_spinner.set()
            spinner_thread.join()
            print("\nError: Unable to setup ChromeDriver automatically")
            print("Please download ChromeDriver manually:")
            print(f"1. Visit: https://googlechromelabs.github.io/chrome-for-testing/")
            print(f"2. Download version matching Chrome {chrome_major}")
            print("3. Extract chromedriver.exe to the same folder as this program")
            return False
            
    except Exception as e:
        if 'stop_spinner' in locals() and not stop_spinner.is_set():
            stop_spinner.set()
            spinner_thread.join()
        print("\nError: Unable to verify Chrome setup")
        return False

def validate_fluxx_url(url):
    """Validate and format Fluxx URL"""
    url = url.strip().lower()
    
    # Remove any protocol prefixes
    if url.startswith(('http://', 'https://')):
        url = url.replace('http://', '').replace('https://', '')
    
    # Remove any trailing slashes
    url = url.rstrip('/')
    
    # Check if URL ends with .fluxx.io
    if not url.endswith('.fluxx.io'):
        if '.fluxx.io' in url:
            # Extract the part before .fluxx.io if it exists
            url = url.split('.fluxx.io')[0] + '.fluxx.io'
        else:
            # Append .fluxx.io if missing
            url = url + '.fluxx.io'
    
    return f'https://{url}'

def wait_for_modal_load(driver, timeout=10):
    """Wait for modal to fully load"""
    try:
        wait = WebDriverWait(driver, timeout)
        modal = wait.until(EC.presence_of_element_located(
            (By.CSS_SELECTOR, "div.modal.new-modal.area[style*='opacity: 1']")
        ))
        
        # Additional wait for content
        wait.until(EC.presence_of_element_located(
            (By.CSS_SELECTOR, "div.modal textarea.code-to-submit")
        ))
        return modal
    except TimeoutException:
        return None

def safely_close_modal(driver, timeout=10):
    """Safely close the modal window"""
    try:
        wait = WebDriverWait(driver, timeout)
        close_button = wait.until(EC.element_to_be_clickable(
            (By.CSS_SELECTOR, "a.close-modal")
        ))
        close_button.click()
        # Wait for modal to disappear
        wait.until(EC.invisibility_of_element_located(
            (By.CSS_SELECTOR, "div.modal.new-modal.area")
        ))
        return True
    except:
        return False

def close_all_models(driver):
    """Close all open models"""
    try:
        # Find all open models
        open_models = driver.find_elements(By.CSS_SELECTOR, "ul.toggle-class.open")
        for model in open_models:
            # Remove the open class
            driver.execute_script("arguments[0].classList.remove('open');", model)
        time.sleep(0.5)  # Brief pause to let animations complete
    except Exception as e:
        print(f"Warning: Could not close all models: {str(e)}")

def ensure_model_open(driver, model_ul, model_name):
    """Ensure a model is open and ready for processing"""
    try:
        # Check if model is already open
        if 'open' not in model_ul.get_attribute('class').split():
            # Find and click the model header
            model_header = model_ul.find_element(By.CSS_SELECTOR, "li.list-label div.link.is-admin")
            driver.execute_script("arguments[0].click();", model_header)
            
            # Wait for the open class to appear
            wait = WebDriverWait(driver, 10)
            wait.until(lambda d: 'open' in model_ul.get_attribute('class').split())
            time.sleep(1)  # Additional pause to let content load
            
            # Verify the model opened
            if 'open' not in model_ul.get_attribute('class').split():
                return False
        return True
    except Exception as e:
        return False

def get_theme_code(driver, theme_element, model_ul, model_name):
    """Get Before/After code for a theme"""
    try:
        # Ensure model is open before processing themes
        if not ensure_model_open(driver, model_ul, model_name):
            return None
        
        # Now ensure the theme is visible by clicking the theme name
        try:
            theme_link = theme_element.find_element(By.CSS_SELECTOR, 
                "a.link.scroll-to-card")
            theme_link.click()
            time.sleep(1)
        except Exception as e:
            return None
            
        # Wait for theme to be visible and expanded
        wait = WebDriverWait(driver, 10)
        wait.until(EC.visibility_of(theme_element))
        
        # Find and click the gear icon
        try:
            gear_icon = theme_element.find_element(By.CSS_SELECTOR, 
                "a.to-modal.open-config[data-on-success='matchListItem,close']")
            driver.execute_script("arguments[0].scrollIntoView(true);", gear_icon)
            time.sleep(1)
            driver.execute_script("arguments[0].click();", gear_icon)
        except Exception as e:
            return None
        
        # Wait for modal to load
        modal = wait_for_modal_load(driver)
        if not modal:
            return None
            
        # Find all code textareas with correct IDs
        code_blocks = {
            'current_before_new': {
                'id': 'model_theme_unsafe_before_new_block',
                'label': 'Current Before New Block'
            },
            'draft_before_new': {
                'id': 'model_theme_draft_before_new_block',
                'label': 'Draft Before New Block'
            },
            'current_after_create': {
                'id': 'model_theme_unsafe_after_create_block',
                'label': 'Current After Create Block'
            },
            'draft_after_create': {
                'id': 'model_theme_draft_after_create_block',
                'label': 'Draft After Create Block'
            }
        }
        
        code_data = {}
        for key, block in code_blocks.items():
            try:
                textarea = modal.find_element(By.CSS_SELECTOR, f"textarea#{block['id']}")
                code = textarea.get_attribute("value").strip()
                code_data[key] = code if code else "N/A"
            except:
                code_data[key] = "N/A"
        
        # Safely close the modal
        safely_close_modal(driver)
        return code_data
        
    except Exception as e:
        try:
            safely_close_modal(driver)
        except:
            pass
        return None

def gather_theme_code(driver, models):
    """Gather Before/After code for all themes"""
    print("\n" + "=" * 80)
    print("\n                     Theme Code Gathering Process")
    print("\n" + "=" * 80 + "\n")
    print("This process will gather Before/After code from all themes automatically.")
    print("This may take several minutes to complete depending on the number of models.")
    print("\nIMPORTANT:")
    print("- You can stop the process at any time by pressing Ctrl+C")
    print("- Please do not interact with the browser while the process is running")
    print("- The browser will automatically handle all interactions")
    print("\n" + "-" * 80)
    verify = input("\nWould you like to proceed with gathering code? (y/n): ").strip().lower()
    
    if verify != 'y':
        print("\nSkipping code gathering process.")
        return models
    
    total_models = len(models)
    current_model = 0
    
    try:
        for model_name, model_data in models.items():
            try:
                current_model += 1
                progress = (current_model / total_models) * 100
                
                # Update progress bar
                bar_length = 50
                filled_length = int(bar_length * current_model // total_models)
                bar = '=' * filled_length + '-' * (bar_length - filled_length)
                sys.stdout.write('\r' + ' ' * 100)  # Clear line
                sys.stdout.write('\rProcessing Models: [{0}] {1:.1f}% ({2}/{3})'.format(
                    bar, progress, current_model, total_models
                ))
                sys.stdout.flush()
                
                model_ul = driver.find_element(By.CSS_SELECTOR, f"ul#{model_name.lower().replace(' ', '_')}")
                
                # Ensure model is open
                if not ensure_model_open(driver, model_ul, model_name):
                    continue
                
                for theme_name, theme_data in model_data['themes'].items():
                    try:
                        theme_elements = model_ul.find_elements(By.CSS_SELECTOR, "li.icon[data-card-uid]")
                        
                        for theme_element in theme_elements:
                            try:
                                label = theme_element.find_element(By.CSS_SELECTOR, 
                                    "a.link.scroll-to-card span.label").text
                                
                                if label == theme_name:
                                    code_data = get_theme_code(driver, theme_element, model_ul, model_name)
                                    if code_data:
                                        models[model_name]['themes'][theme_name]['code'] = code_data
                            except Exception:
                                continue
                    except Exception:
                        continue
                
                # Close model after processing
                try:
                    if 'open' in model_ul.get_attribute('class').split():
                        model_header = model_ul.find_element(By.CSS_SELECTOR, "li.list-label div.link.is-admin")
                        driver.execute_script("arguments[0].click();", model_header)
                        time.sleep(1)
                except Exception:
                    pass
                    
            except Exception:
                continue
        
        # Show completion
        sys.stdout.write('\r' + ' ' * 100)  # Clear line
        sys.stdout.write('\rCode gathering process complete!')
        sys.stdout.flush()
        print("\n" + "=" * 80)
        
        return models
        
    except Exception as e:
        print(f"\nError during code gathering: {str(e)}")
        return models

def main():
    try:
        # Show logo and contact info
        print_logo()
        
        # Get Fluxx URL
        url = get_fluxx_url()
        if not url:
            return

        # Check Chrome setup
        if not check_chrome_and_driver():
            input("\nPress Enter to exit...")
            return
            
        # Setup Chrome
        print("Starting Chrome...")
        driver_path = get_resource_path("chromedriver.exe")
        options = webdriver.ChromeOptions()
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--log-level=3')
        options.add_experimental_option('excludeSwitches', ['enable-logging', 'enable-automation'])
        options.add_experimental_option('useAutomationExtension', False)
        
        # Create temp profile
        temp_dir = os.path.join(os.getcwd(), 'chrome_temp')
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
        options.add_argument(f'--user-data-dir={temp_dir}')
        
        service = Service(driver_path, log_path='NUL')  # Suppress ChromeDriver logs
        driver = webdriver.Chrome(service=service, options=options)
        
        print(f"Navigating to {url}")
        driver.get(url)
        
        current_url = driver.current_url
        print(f"Current URL: {current_url}\n")  # Add newline after URL
        
        if current_url == "data:,":
            print("Navigation failed. Please check your internet connection.")
            return
        
        # Wait for dashboard and navigate to Admin Panel
        if not wait_for_dashboard(driver):
            print("\nError: Could not detect dashboard load.")
            input("Press Enter to exit...")
            return
        
        print("\nDashboard detected! Navigating to Admin Panel...")
        
        # Navigate to Admin Panel
        if not navigate_to_admin(driver):
            print("\nError: Could not navigate to Admin Panel.")
            input("Press Enter to exit...")
            return
        
        # Parse the Forms section
        while True:  # Options loop
            # Get the data
            models_data = wait_for_forms_and_parse(driver)
            if not models_data:
                print("\nError: Could not parse Forms section.")
                print_divider()
                retry_choice = input("Would you like to return to the main menu? (y/n): ").strip().lower()
                if retry_choice == 'y':
                    continue  # Continue to retry
                else:
                    input("Press Enter to exit...")
                    return
            
            # Ask if user wants to gather code blocks
            print("\nWould you like to gather Before/After code blocks from themes?")
            print("This step can be skipped if you only need model structure and workflows.")
            code_choice = input("Gather code blocks? (y/n): ").strip().lower()
            
            if code_choice == 'y':
                # Gather theme code if requested
                models_data = gather_theme_code(driver, models_data)
            
            # Ask if user wants to scan methods
            print("\nWould you like to scan model methods?")
            print("This will gather methods from each model's themes.")
            methods_choice = input("Scan methods? (y/n): ").strip().lower()
            
            if methods_choice == 'y':
                # Scan methods
                models_data = scan_methods(driver, models_data)
            
            # Ask if user wants to scan workflows
            print("\nWould you like to scan model workflows?")
            print("This will gather workflow states and actions for each model.")
            workflow_choice = input("Scan workflows? (y/n): ").strip().lower()
            
            if workflow_choice == 'y':
                # Scan workflows
                models_data = scan_model_workflows(driver, models_data)

            print_divider()
            print("Available Actions:")
            print("1. Generate Word document")
            print("2. Re-run scan")
            print("3. Exit")
            
            choice = input("\nEnter your choice (1-3): ").strip()
            
            if choice == '1':
                doc_filename = wait_with_spinner(
                    "Generating Word document...", 
                    generate_word_document, 
                    models_data,
                    site_url=url  # Pass the URL to the document generator
                )
                if doc_filename:
                    print(f"\nDocumentation has been saved to: {doc_filename}")
                
                print_divider()
                print("Would you like to:")
                print("1. Return to options")
                print("2. Exit")
                sub_choice = input("\nEnter your choice (1-2): ").strip()
                if sub_choice == '2':
                    return
                    
            elif choice == '2':
                print("\nRe-running scan...")
                continue
                
            elif choice == '3':
                return
                
            else:
                print("\nInvalid choice. Please try again.")
                continue
        
    except KeyboardInterrupt:
        print("\n\nProcess interrupted by user.")
    except Exception as e:
        print("\nAn error occurred. Details below:")
        print("=" * 50)
        print(f"Error type: {type(e).__name__}")
        print(f"Error message: {str(e)}")
        import traceback
        print("\nFull error traceback:")
        print("-" * 50)
        traceback.print_exc()
        print("=" * 50)
    finally:
        try:
            driver.quit()
            # Clean up temp directory
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception:
                pass
        except Exception:
            pass

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\nScript terminated by user.")
    except Exception as e:
        print("\nAn error occurred. Details below:")
        print("=" * 50)
        print(f"Error type: {type(e).__name__}")
        print(f"Error message: {str(e)}")
        import traceback
        print("\nFull error traceback:")
        print("-" * 50)
        traceback.print_exc()
        print("=" * 50)
    finally:
        print("\nPress Enter to exit...")
        input()  # This will keep the window open

